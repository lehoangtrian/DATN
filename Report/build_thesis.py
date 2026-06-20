# -*- coding: utf-8 -*-
"""Sinh file khoa luan tot nghiep PhoneStore (.docx) bang python-docx."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"D:\0_DATN\DATN"
DIAGRAMS = os.path.join(BASE, "Report", "diagrams")
SHOTS = os.path.join(BASE, "Report", "screenshots")
OUT = os.path.join(BASE, "Report", "KhoaLuanTotNghiep_PhoneStore.docx")

FONT = "Times New Roman"
SIZE = 13

# ===================== DANH SACH HINH / BANG (nguon duy nhat) =====================
FIGURES = {
    "1.1": ("Sơ đồ kiến trúc hệ thống PhoneStore", DIAGRAMS, "01_kien_truc_he_thong.png"),
    "3.1": ("Use Case Diagram - Tổng quan hệ thống PhoneStore", DIAGRAMS, "02_use_case_diagram.png"),
    "4.1": ("Class Diagram - Tổng quan toàn hệ thống", DIAGRAMS, "03_class_diagram_tong_quan.png"),
    "4.2": ("Class Diagram - Nhóm Sản phẩm và Mua hàng", DIAGRAMS, "04_class_diagram_san_pham_mua_hang.png"),
    "4.3": ("Class Diagram - Nhóm Khuyến mãi và Tài chính", DIAGRAMS, "05_class_diagram_khuyen_mai_tai_chinh.png"),
    "4.4": ("Class Diagram - Nhóm Hỗ trợ và Tương tác khách hàng", DIAGRAMS, "06_class_diagram_ho_tro_tuong_tac.png"),
    "4.5": ("Sequence Diagram - Đặt hàng (Checkout)", DIAGRAMS, "07_sequence_dat_hang.png"),
    "4.6": ("Collaboration Diagram - Đặt hàng", DIAGRAMS, "17_collaboration_diagram_dat_hang.png"),
    "4.7": ("Sequence Diagram - Giao hàng thành công và tích điểm thưởng", DIAGRAMS, "08_sequence_giao_hang_tich_diem.png"),
    "4.8": ("Sequence Diagram - Yêu cầu trả hàng và hoàn tiền", DIAGRAMS, "09_sequence_tra_hang_hoan_tien.png"),
    "4.9": ("Sequence Diagram - Đổi điểm thưởng sang mã giảm giá", DIAGRAMS, "10_sequence_doi_diem_giam_gia.png"),
    "4.10": ("Sequence Diagram - Chatbot AI trả lời câu hỏi khách hàng (RAG)", DIAGRAMS, "11_sequence_chatbot_rag.png"),
    "4.11": ("Activity Diagram - Luồng đặt hàng (góc nhìn khách hàng)", DIAGRAMS, "12_activity_dat_hang.png"),
    "4.12": ("State Diagram - Vòng đời đơn hàng (Order)", DIAGRAMS, "16_state_diagram_order.png"),
    "4.13": ("Package Diagram - Tổ chức module hệ thống", DIAGRAMS, "13_package_diagram.png"),
    "4.14": ("Component Diagram - Các thành phần logic của hệ thống", DIAGRAMS, "14_component_diagram.png"),
    "4.15": ("Deployment Diagram - Sơ đồ triển khai vật lý", DIAGRAMS, "15_deployment_diagram.png"),
    "5.1": ("Trang chủ website PhoneStore", SHOTS, "01_home.png"),
    "5.2": ("Trang danh sách sản phẩm", SHOTS, "02_product_list.png"),
    "5.3": ("Trang chi tiết sản phẩm", SHOTS, "03_product_detail.png"),
    "5.4": ("Trang đăng nhập", SHOTS, "04_login.png"),
    "5.5": ("Trang giỏ hàng", SHOTS, "05_cart.png"),
    "5.6": ("Trang thanh toán (checkout)", SHOTS, "06_checkout.png"),
    "5.7": ("Trang hồ sơ cá nhân và điểm thưởng", SHOTS, "07_profile.png"),
    "5.8": ("Trang danh sách đơn hàng của khách hàng", SHOTS, "08_orders.png"),
    "5.9": ("Khung chat hỗ trợ với Chatbot AI", SHOTS, "09_chatbot.png"),
    "5.10": ("Trang Dashboard quản trị", SHOTS, "10_admin_dashboard.png"),
    "5.11": ("Trang quản lý sản phẩm (Admin)", SHOTS, "11_admin_products.png"),
    "5.12": ("Trang quản lý đơn hàng (Admin)", SHOTS, "12_admin_orders.png"),
}

TABLES_INDEX = [
    ("2.1", "Ánh xạ các giai đoạn RUP vào hoạt động thực hiện đề tài"),
    ("3.1", "Danh sách Actor trong hệ thống"),
    ("3.2", "Tổng hợp các Use Case theo nhóm chức năng"),
    ("3.3", "Đặc tả Use Case: Đăng ký / Đăng nhập tài khoản"),
    ("3.4", "Đặc tả Use Case: Đặt hàng (Checkout)"),
    ("3.5", "Đặc tả Use Case: Áp dụng mã giảm giá"),
    ("3.6", "Đặc tả Use Case: Đổi điểm thưởng sang mã giảm giá"),
    ("3.7", "Đặc tả Use Case: Yêu cầu trả hàng"),
    ("3.8", "Đặc tả Use Case: Chat với hệ thống (Chatbot AI)"),
    ("3.9", "Đặc tả Use Case: Quản lý sản phẩm (Admin)"),
    ("3.10", "Đặc tả Use Case: Quản lý đơn hàng (Admin)"),
    ("4.1", "Danh sách collection dữ liệu trong MongoDB"),
    ("5.1", "Công nghệ sử dụng phía Client"),
    ("5.2", "Công nghệ sử dụng phía Server"),
    ("5.3", "Công nghệ sử dụng cho Chatbot AI"),
    ("5.4", "Thống kê API đã kiểm thử bằng Postman"),
]

# ===================== HELPER =====================

def set_run_font(run, size=SIZE, bold=False, italic=False, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)


def add_para(doc, text="", size=SIZE, bold=False, italic=False, align='justify',
             space_after=6, space_before=0, line_spacing=1.5, indent_first=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    align_map = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT, 'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items, size=SIZE):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(it)
        set_run_font(run, size=size)


def add_heading(doc, text, level=1, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    sizes = {1: 17, 2: 14, 3: 13}
    set_run_font(run, size=sizes.get(level, 13), bold=True)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 8)
    p.paragraph_format.keep_with_next = True
    return p


def add_figure(doc, fig_no, width_cm=15.2):
    caption, folder, filename = FIGURES[fig_no]
    path = os.path.join(folder, filename)
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(f"Hình {fig_no}. {caption}")
    set_run_font(run, size=12, bold=True)


def add_table_caption(doc, table_no):
    caption = dict(TABLES_INDEX)[table_no]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Bảng {table_no}. {caption}")
    set_run_font(run, size=12, bold=True)


def shade_cell(cell, color="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=12, align='left'):
    cell.text = ""
    p = cell.paragraphs[0]
    align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=12, align='center')
        shade_cell(cell)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=12)
    if widths:
        for i, w in enumerate(widths):
            for r in table.rows:
                r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return table


def add_usecase_table(doc, table_no, name, short_desc, actor, pre, post, main_flow, alt_flows, exception_flows):
    add_table_caption(doc, table_no)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def kv_row(label, content_lines):
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True, size=12)
        shade_cell(row[0])
        row[0].width = Cm(3.8)
        cell = row[1]
        cell.text = ""
        if isinstance(content_lines, str):
            content_lines = [content_lines]
        if not content_lines:
            content_lines = ["Không có."]
        for i, line in enumerate(content_lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(line)
            set_run_font(run, size=12)
        row[1].width = Cm(11.4)

    kv_row("Tên Use Case", name)
    kv_row("Actor liên quan", actor)
    kv_row("Mô tả ngắn", short_desc)
    kv_row("Điều kiện trước\n(Pre-condition)", pre)
    kv_row("Điều kiện sau\n(Post-condition)", post)
    kv_row("Luồng chính\n(Main Flow)", main_flow)
    kv_row("Luồng phụ\n(Alternate Flow)", alt_flows)
    kv_row("Luồng ngoại lệ\n(Exception Flow)", exception_flows)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def add_toc_field(doc, switch='TOC \\o "1-3" \\h \\z \\u'):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = switch
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    txt = OxmlElement('w:t')
    txt.text = "Nhấp phải vào đây và chọn \"Update Field\" (hoặc nhấn F9) trong Word để tự động cập nhật mục lục."
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(txt); r.append(fld_end)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_end)
    set_run_font(run, size=11)


def set_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


# ===================== KHOI TAO DOCUMENT =====================
doc = Document()
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(SIZE)
set_margins(doc)
add_page_number_footer(doc)

print("Doc khoi tao xong, bat dau ghi noi dung...")

# ===================== TRANG BIA =====================
for _ in range(2):
    add_para(doc, "")
p = add_para(doc, "TRƯỜNG ĐẠI HỌC ...........................", align='center', size=14, bold=True)
add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN", align='center', size=14, bold=True)
add_para(doc, "")
add_para(doc, "")
add_para(doc, "KHÓA LUẬN TỐT NGHIỆP", align='center', size=26, bold=True)
add_para(doc, "")
run_p = add_para(doc, "XÂY DỰNG HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ BÁN ĐIỆN THOẠI",
                  align='center', size=18, bold=True)
add_para(doc, "TÍCH HỢP CHATBOT TƯ VẤN BẰNG TRÍ TUỆ NHÂN TẠO (PHONESTORE)",
         align='center', size=18, bold=True)
add_para(doc, "")
add_para(doc, "")
add_para(doc, "")
add_para(doc, "Sinh viên thực hiện : [Họ và tên sinh viên]", align='center', size=14)
add_para(doc, "Mã số sinh viên     : [MSSV]", align='center', size=14)
add_para(doc, "Lớp                 : [Lớp]", align='center', size=14)
add_para(doc, "Giảng viên hướng dẫn : [Họ và tên GVHD]", align='center', size=14)
add_para(doc, "")
add_para(doc, "")
add_para(doc, "Năm 2026", align='center', size=14, bold=True)

# ===================== LOI CAM ON =====================
add_heading(doc, "LỜI CẢM ƠN", level=1, page_break_before=True)
add_para(doc,
    "Trong suốt quá trình học tập, nghiên cứu và thực hiện khóa luận tốt nghiệp với đề tài "
    "\"Xây dựng hệ thống thương mại điện tử bán điện thoại tích hợp chatbot tư vấn bằng trí tuệ "
    "nhân tạo (PhoneStore)\", em đã nhận được rất nhiều sự giúp đỡ, hướng dẫn và tạo điều kiện từ "
    "quý thầy cô trong Khoa Công nghệ thông tin.", indent_first=1)
add_para(doc,
    "Em xin gửi lời cảm ơn chân thành đến giảng viên hướng dẫn đã tận tình chỉ bảo, góp ý và "
    "định hướng cho em trong suốt thời gian thực hiện đề tài, giúp em hoàn thiện cả về kiến thức "
    "chuyên môn lẫn phương pháp nghiên cứu khoa học.", indent_first=1)
add_para(doc,
    "Em cũng xin cảm ơn quý thầy cô đã truyền đạt những kiến thức nền tảng quý báu trong suốt "
    "quá trình học tập tại trường, đặc biệt là các kiến thức về phân tích thiết kế hướng đối tượng, "
    "công nghệ web và trí tuệ nhân tạo — những nền tảng quan trọng giúp em hoàn thành đề tài này.", indent_first=1)
add_para(doc,
    "Do thời gian thực hiện và kiến thức còn hạn chế, khóa luận chắc chắn không tránh khỏi những "
    "thiếu sót. Em rất mong nhận được sự góp ý của quý thầy cô để đề tài được hoàn thiện hơn.", indent_first=1)
add_para(doc, "Em xin chân thành cảm ơn!", indent_first=1)
add_para(doc, "")
add_para(doc, "Tp. Hồ Chí Minh, tháng 6 năm 2026", align='right', italic=True)
add_para(doc, "Sinh viên thực hiện", align='right', italic=True)

# ===================== MUC LUC =====================
add_heading(doc, "MỤC LỤC", level=1, page_break_before=True)
add_toc_field(doc)

# ===================== DANH MUC TU VIET TAT =====================
add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", level=1, page_break_before=True)
abbr_rows = [
    ("UML", "Unified Modeling Language - Ngôn ngữ mô hình hóa thống nhất"),
    ("OOAD", "Object-Oriented Analysis and Design - Phân tích & thiết kế hướng đối tượng"),
    ("RUP", "Rational Unified Process - Quy trình phát triển hợp nhất"),
    ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
    ("REST", "Representational State Transfer - Kiến trúc API dựa trên HTTP"),
    ("JWT", "JSON Web Token - Chuẩn token xác thực người dùng"),
    ("OTP", "One-Time Password - Mã xác thực dùng một lần"),
    ("SPA", "Single Page Application - Ứng dụng web một trang"),
    ("CRUD", "Create, Read, Update, Delete - Bốn tác vụ cơ bản trên dữ liệu"),
    ("LLM", "Large Language Model - Mô hình ngôn ngữ lớn"),
    ("RAG", "Retrieval-Augmented Generation - Sinh văn bản có truy xuất dữ liệu thật"),
    ("IPN", "Instant Payment Notification - Thông báo thanh toán tức thời (VNPay)"),
    ("CSDL", "Cơ sở dữ liệu"),
    ("VNPay", "Cổng thanh toán trực tuyến VNPay"),
]
add_simple_table(doc, ["Từ viết tắt", "Diễn giải"], abbr_rows, widths=[3, 12.2])

# ===================== DANH MUC HINH ANH =====================
add_heading(doc, "DANH MỤC HÌNH ẢNH", level=1, page_break_before=True)
for fig_no, (caption, _, _) in FIGURES.items():
    add_para(doc, f"Hình {fig_no}. {caption}", size=12, align='left', space_after=4)

# ===================== DANH MUC BANG =====================
add_heading(doc, "DANH MỤC BẢNG", level=1, page_break_before=True)
for table_no, caption in TABLES_INDEX:
    add_para(doc, f"Bảng {table_no}. {caption}", size=12, align='left', space_after=4)

# ===================== CHUONG 1 =====================
add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", level=1, page_break_before=True)

add_heading(doc, "1.1. Lý do chọn đề tài", level=2)
add_para(doc,
    "Cùng với sự phát triển mạnh mẽ của Internet và thiết bị di động, thương mại điện tử đã trở "
    "thành một trong những lĩnh vực tăng trưởng nhanh nhất hiện nay, trong đó ngành hàng điện "
    "thoại di động luôn nằm trong nhóm sản phẩm được giao dịch trực tuyến nhiều nhất do đặc tính "
    "giá trị cao, thông số kỹ thuật rõ ràng và nhu cầu so sánh, tư vấn trước khi mua rất lớn. Việc "
    "xây dựng một hệ thống bán hàng trực tuyến chuyên biệt cho ngành hàng điện thoại, có đầy đủ các "
    "nghiệp vụ thực tế như giỏ hàng, khuyến mãi, điểm thưởng, đổi trả... là một bài toán vừa mang "
    "tính học thuật vừa có giá trị ứng dụng cao.", indent_first=1)
add_para(doc,
    "Bên cạnh đó, một trong những hạn chế lớn của các website bán hàng truyền thống là khách hàng "
    "thường phải tự tìm hiểu thông tin sản phẩm hoặc chờ đợi nhân viên tư vấn, gây trải nghiệm chưa "
    "tốt và tốn nhân lực vận hành. Sự phát triển của các mô hình ngôn ngữ lớn (Large Language Model "
    "- LLM) mở ra cơ hội xây dựng chatbot tư vấn tự động, hoạt động 24/7, trả lời dựa trên dữ liệu "
    "thật của hệ thống (giá, tồn kho, khuyến mãi) thay vì chỉ dựa vào kịch bản trả lời cố định.", indent_first=1)
add_para(doc,
    "Về mặt học thuật, đây là cơ hội để vận dụng một cách có hệ thống phương pháp Phân tích và "
    "Thiết kế Hướng đối tượng (Object-Oriented Analysis and Design - OOAD) cùng ngôn ngữ mô hình hóa "
    "UML vào một dự án phần mềm thực tế có quy mô vừa, từ giai đoạn xác định yêu cầu, thiết kế, đến "
    "cài đặt và kiểm thử — qua đó củng cố và minh chứng năng lực phân tích thiết kế hệ thống của "
    "sinh viên. Xuất phát từ những lý do trên, em chọn đề tài \"Xây dựng hệ thống thương mại điện tử "
    "bán điện thoại tích hợp chatbot tư vấn bằng trí tuệ nhân tạo (PhoneStore)\" làm đề tài khóa luận "
    "tốt nghiệp.", indent_first=1)
add_para(doc,
    "Hình 1.1 giới thiệu tổng quan kiến trúc của hệ thống PhoneStore sẽ được xây dựng trong khóa "
    "luận, gồm ba tầng (giao diện, ứng dụng, dữ liệu) cùng dịch vụ chatbot AI và các tác vụ định kỳ "
    "hỗ trợ vận hành; kiến trúc này được phân tích chi tiết tại Chương 4.", indent_first=1)
add_figure(doc, "1.1")

add_heading(doc, "1.2. Mục tiêu đề tài", level=2)
add_para(doc, "Đề tài hướng tới các mục tiêu cụ thể sau:", indent_first=1)
add_bullets(doc, [
    "Phân tích và thiết kế hệ thống theo phương pháp OOAD/UML một cách bài bản, bám sát quy trình "
    "lặp và tăng trưởng (Iterative, Incremental Framework) được trình bày trong tài liệu tham khảo [1].",
    "Xây dựng hoàn chỉnh một hệ thống thương mại điện tử bán điện thoại với đầy đủ các nghiệp vụ: "
    "quản lý sản phẩm theo biến thể (màu sắc, dung lượng), giỏ hàng, đặt hàng, thanh toán (COD, "
    "ví điện tử, VNPay), khuyến mãi (Flash Sale, mã giảm giá), điểm thưởng theo hạng thành viên, "
    "đánh giá sản phẩm, đổi trả hàng và quản trị vận hành.",
    "Tích hợp một chatbot tư vấn bằng AI sử dụng mô hình ngôn ngữ lớn đã được tinh chỉnh (fine-tune) "
    "riêng cho lĩnh vực bán điện thoại, hoạt động theo cơ chế truy xuất dữ liệu thật trước khi sinh "
    "câu trả lời (Retrieval-Augmented Generation - RAG) để đảm bảo thông tin tư vấn (giá, tồn kho) "
    "luôn chính xác, không suy diễn.",
    "Cài đặt hệ thống bằng các công nghệ web hiện đại (React, Node.js/Express, MongoDB) kết hợp với "
    "dịch vụ AI viết bằng Python (FastAPI), và kiểm thử hệ thống trên dữ liệu, môi trường thực tế.",
])

add_heading(doc, "1.3. Đối tượng và phạm vi nghiên cứu", level=2)
add_para(doc, "Đối tượng nghiên cứu của đề tài bao gồm:", indent_first=1)
add_bullets(doc, [
    "Khách hàng (Guest/Customer): người dùng cuối có nhu cầu tìm hiểu, so sánh và mua điện thoại "
    "trực tuyến.",
    "Quản trị viên/Nhân viên (Admin/Staff): người vận hành, quản lý sản phẩm, đơn hàng, khuyến mãi "
    "và hỗ trợ khách hàng.",
    "Quy trình nghiệp vụ bán lẻ điện thoại trực tuyến: từ tìm kiếm sản phẩm, đặt hàng, thanh toán, "
    "giao hàng, tích điểm, đến đổi trả và hậu mãi.",
])
add_para(doc, "Phạm vi thực hiện của đề tài:", indent_first=1)
add_bullets(doc, [
    "Xây dựng đầy đủ phần phía khách hàng (storefront) và phần quản trị (admin dashboard) trên nền "
    "web, chạy được trên môi trường máy chủ cục bộ (local).",
    "Chatbot AI được huấn luyện tinh chỉnh và chạy bằng công cụ Ollama trên máy chủ AI cục bộ; cổng "
    "thanh toán VNPay được tích hợp ở môi trường thử nghiệm (sandbox), chưa triển khai thanh toán "
    "thật.",
    "Đề tài không bao gồm việc xây dựng ứng dụng di động riêng (native mobile app) và chưa thực hiện "
    "kiểm thử tải (load test) ở quy mô sản xuất thực tế.",
])

add_heading(doc, "1.4. Phương pháp nghiên cứu", level=2)
add_para(doc,
    "Khóa luận áp dụng phương pháp Phân tích và Thiết kế Hướng đối tượng (OOAD) với ngôn ngữ mô hình "
    "hóa UML làm phương pháp luận chủ đạo, theo khung tiến trình lặp và tăng trưởng (Iterative, "
    "Incremental Framework) gồm bốn giai đoạn Khởi tạo (Inception) - Phân tích/Thiết kế chi tiết "
    "(Elaboration) - Xây dựng (Construction) - Chuyển giao (Transition), được trình bày trong tài "
    "liệu \"UML Applied - Object Oriented Analysis and Design using the UML\" [1]. Nội dung và cách "
    "áp dụng cụ thể từng giai đoạn cho đề tài được trình bày chi tiết tại Chương 2.", indent_first=1)
add_para(doc,
    "Toàn bộ các sơ đồ UML sử dụng trong khóa luận — gồm Use Case Diagram, Class Diagram, Sequence "
    "Diagram, Collaboration Diagram, Activity Diagram, State Diagram, Package Diagram, Component "
    "Diagram và Deployment Diagram — được xây dựng bám theo đúng cú pháp và mục đích sử dụng của "
    "từng loại sơ đồ theo tài liệu [1], dựa trên công cụ PlantUML để đảm bảo tính nhất quán và dễ "
    "kiểm soát phiên bản.", indent_first=1)
add_para(doc,
    "Bên cạnh phương pháp luận thiết kế, đề tài còn sử dụng phương pháp nghiên cứu tài liệu kỹ thuật "
    "(đọc tài liệu chính thức của các công nghệ sử dụng) và phương pháp thực nghiệm (cài đặt, vận "
    "hành thử và kiểm thử hệ thống thật trên dữ liệu mô phỏng) để hoàn thiện sản phẩm.", indent_first=1)

add_heading(doc, "1.5. Cấu trúc báo cáo", level=2)
add_para(doc, "Ngoài phần mở đầu, danh mục và phụ lục, nội dung khóa luận được trình bày trong 6 "
    "chương:", indent_first=1)
add_bullets(doc, [
    "Chương 1 - Tổng quan đề tài: trình bày lý do chọn đề tài, mục tiêu, đối tượng, phạm vi và "
    "phương pháp nghiên cứu.",
    "Chương 2 - Cơ sở lý thuyết: trình bày tổng quan về UML, các loại sơ đồ UML sử dụng trong đề "
    "tài, quy trình phát triển RUP và các nền tảng công nghệ liên quan.",
    "Chương 3 - Phân tích yêu cầu hệ thống: xác định actor, xây dựng Use Case Diagram, đặc tả các "
    "Use Case chính và yêu cầu phi chức năng.",
    "Chương 4 - Thiết kế hệ thống: trình bày mô hình hóa khái niệm (Class Diagram), thiết kế tương "
    "tác đối tượng (Sequence/Collaboration Diagram), thiết kế hành vi (State Diagram), thiết kế "
    "kiến trúc (Package/Component/Deployment Diagram) và thiết kế cơ sở dữ liệu.",
    "Chương 5 - Cài đặt và kiểm thử: trình bày công nghệ sử dụng thực tế, tổ chức mã nguồn, giao "
    "diện hệ thống đã cài đặt và kết quả kiểm thử.",
    "Chương 6 - Kết luận và hướng phát triển: tổng kết kết quả đạt được, hạn chế và hướng phát "
    "triển của đề tài.",
])

# ===================== CHUONG 2 =====================
add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", level=1, page_break_before=True)

add_heading(doc, "2.1. Tổng quan về UML", level=2)
add_para(doc,
    "UML (Unified Modeling Language) là một ngôn ngữ mô hình hóa trực quan dùng để đặc tả, xây "
    "dựng và lập tài liệu cho các thành phần của một hệ thống phần mềm hướng đối tượng. Theo tài "
    "liệu [1], các đối tượng tham gia vào một dự án phần mềm (nhà phân tích, người thiết kế, lập "
    "trình viên, kiểm thử viên, khách hàng...) thường cần những góc nhìn khác nhau và mức độ chi "
    "tiết khác nhau về cùng một hệ thống; UML cung cấp một bộ ký pháp đủ phong phú để mỗi nhóm "
    "đối tượng đều có ít nhất một loại sơ đồ phù hợp với nhu cầu của mình.", indent_first=1)
add_para(doc,
    "Tài liệu [1] tổng kết vai trò của từng loại sơ đồ UML chính bằng một câu hỏi thiết kế mà sơ "
    "đồ đó trả lời, đây cũng chính là cách đề tài này lựa chọn và sử dụng từng loại sơ đồ trong "
    "quá trình phân tích, thiết kế hệ thống PhoneStore:", indent_first=1)
add_bullets(doc, [
    "Use Case Diagram - Hệ thống sẽ tương tác với thế giới bên ngoài như thế nào?",
    "Class Diagram - Hệ thống cần những đối tượng nào? Chúng quan hệ với nhau ra sao?",
    "Collaboration Diagram - Các đối tượng sẽ tương tác với nhau như thế nào?",
    "Sequence Diagram - Các đối tượng sẽ tương tác với nhau như thế nào theo trình tự thời gian?",
    "State Diagram - Đối tượng sẽ ở những trạng thái nào trong vòng đời của nó?",
    "Package Diagram - Hệ thống sẽ được phân chia thành các module như thế nào?",
    "Component Diagram - Các thành phần phần mềm sẽ quan hệ với nhau như thế nào?",
    "Deployment Diagram - Phần mềm sẽ được triển khai như thế nào trên hạ tầng vật lý?",
])

add_heading(doc, "2.2. Các loại sơ đồ UML sử dụng trong khóa luận", level=2)

add_heading(doc, "2.2.1. Use Case Diagram", level=3)
add_para(doc,
    "Theo [1], một Use Case là mô tả hành vi của hệ thống dưới góc nhìn của người dùng. Đây là "
    "sơ đồ có cú pháp đơn giản, dễ hiểu đối với cả người phát triển và khách hàng, vì vậy Use Case "
    "thường được dùng làm điểm khởi đầu xuyên suốt cả quá trình phát triển, từ giai đoạn khởi tạo "
    "cho đến khi triển khai. Trong đề tài, Use Case Diagram được dùng ở Chương 3 để xác định toàn "
    "bộ phạm vi chức năng của hệ thống PhoneStore theo từng nhóm actor.", indent_first=1)

add_heading(doc, "2.2.2. Class Diagram", level=3)
add_para(doc,
    "Class Diagram là một phần thiết yếu của bất kỳ phương pháp thiết kế hướng đối tượng nào. "
    "Theo [1], Class Diagram có thể được sử dụng ngay từ giai đoạn phân tích để vẽ Mô hình khái "
    "niệm (Conceptual Model) — tức các khái niệm chính mà khách hàng/nghiệp vụ quan tâm — cũng như "
    "ở giai đoạn thiết kế chi tiết (Design Class Diagram). Trong đề tài, Class Diagram được dùng ở "
    "cả hai mức: một sơ đồ tổng quan toàn hệ thống và ba sơ đồ chi tiết theo từng nhóm nghiệp vụ "
    "(Chương 4).", indent_first=1)

add_heading(doc, "2.2.3. Collaboration Diagram", level=3)
add_para(doc,
    "Vì phần mềm hướng đối tượng vận hành dựa trên sự cộng tác giữa các đối tượng, Collaboration "
    "Diagram theo [1] được dùng để mô tả cách các đối tượng cộng tác với nhau để hoàn thành một "
    "hành vi cụ thể của hệ thống, thông qua các đối tượng và các thông điệp (message) được đánh số "
    "thứ tự giữa chúng. Đề tài sử dụng Collaboration Diagram để minh họa luồng cộng tác đối tượng "
    "khi xử lý nghiệp vụ đặt hàng.", indent_first=1)

add_heading(doc, "2.2.4. Sequence Diagram", level=3)
add_para(doc,
    "Theo [1], Sequence Diagram thể hiện cùng loại thông tin với Collaboration Diagram nhưng theo "
    "một hình thức khác — các đường kẻ dọc (lifeline) biểu diễn trục thời gian, cho phép quan sát "
    "rõ trình tự trao đổi thông điệp giữa các đối tượng theo thời gian. Đây là loại sơ đồ được sử "
    "dụng nhiều nhất trong Chương 4 của khóa luận để mô tả chi tiết năm luồng nghiệp vụ chính: đặt "
    "hàng, giao hàng và tích điểm, trả hàng và hoàn tiền, đổi điểm thưởng, và chatbot AI trả lời "
    "câu hỏi khách hàng.", indent_first=1)

add_heading(doc, "2.2.5. Activity Diagram", level=3)
add_para(doc,
    "Activity Diagram dùng để mô tả luồng xử lý của một quy trình nghiệp vụ với các điểm rẽ nhánh "
    "(quyết định) và song song, tương tự lưu đồ (flowchart) nhưng theo ký pháp UML. Đề tài sử dụng "
    "Activity Diagram để mô tả toàn bộ luồng đặt hàng dưới góc nhìn của khách hàng, bao gồm các "
    "điểm rẽ nhánh quan trọng như kiểm tra mã giảm giá, điểm thưởng và tồn kho.", indent_first=1)

add_heading(doc, "2.2.6. State Diagram", level=3)
add_para(doc,
    "Theo [1], nhiều đối tượng trong hệ thống tại một thời điểm sẽ ở một trạng thái nhất định, và "
    "việc không kiểm soát chặt chẽ các phép chuyển trạng thái (state transition) có thể gây ra "
    "những lỗi nghiêm trọng cho hệ thống. State Diagram cung cấp ký pháp để mô hình hóa các trạng "
    "thái hợp lệ và các phép chuyển giữa chúng. Trong đề tài, State Diagram được dùng để mô tả "
    "vòng đời của đối tượng Order — từ khi đặt hàng đến khi giao thành công, hủy hoặc trả hàng.", indent_first=1)

add_heading(doc, "2.2.7. Package Diagram", level=3)
add_para(doc,
    "Theo [1], một hệ thống không tầm thường nào cũng cần được phân chia thành các phần nhỏ hơn, "
    "dễ hiểu hơn (\"chunks\"), và Package Diagram cho phép mô hình hóa việc phân chia này một cách "
    "đơn giản và hiệu quả. Đề tài sử dụng Package Diagram để thể hiện cách tổ chức các module mã "
    "nguồn của ba thành phần Client, Server và Bot.", indent_first=1)

add_heading(doc, "2.2.8. Component Diagram và Deployment Diagram", level=3)
add_para(doc,
    "Theo [1], Component Diagram có chức năng tương tự Package Diagram nhưng nhấn mạnh vào các "
    "thành phần phần mềm vật lý (file, thư viện liên kết, file thực thi...) và sự phụ thuộc giữa "
    "chúng, thay vì sự phân chia logic. Deployment Diagram đi xa hơn một bước, cho phép lập kế "
    "hoạch về cách phần mềm sẽ được triển khai trên hạ tầng phần cứng thực tế. Đề tài sử dụng hai "
    "loại sơ đồ này để mô tả các thành phần dịch vụ (Auth, Product, Order, Payment, Loyalty, "
    "Notification, Chatbot) và cách chúng được triển khai trên các máy chủ (Application Server, "
    "AI Server, Database Server).", indent_first=1)

add_heading(doc, "2.3. Quy trình phát triển: Khung tiến trình lặp và tăng trưởng", level=2)
add_para(doc,
    "Theo [1], một khung tiến trình lặp và tăng trưởng (Iterative, Incremental Framework) khắc "
    "phục được nhiều nhược điểm của mô hình thác nước (waterfall) truyền thống, và được chia thành "
    "bốn giai đoạn chính: Khởi tạo (Inception), Phân tích/Thiết kế chi tiết (Elaboration), Xây dựng "
    "(Construction) và Chuyển giao (Transition). Bốn giai đoạn này được thực hiện tuần tự nhưng "
    "không nên bị nhầm lẫn với các pha của mô hình thác nước, vì mỗi giai đoạn đều có thể bao gồm "
    "nhiều vòng lặp phân tích - thiết kế - cài đặt - kiểm thử nhỏ bên trong nó.", indent_first=1)
add_para(doc,
    "Tài liệu [1] cũng cho biết Rational Unified Process (RUP) — do cùng nhóm tác giả phát triển "
    "UML xây dựng — là ví dụ tiêu biểu nhất của một khung tiến trình lặp và tăng trưởng, và khóa "
    "luận này lấy phần cốt lõi của RUP (bốn giai đoạn nêu trên) làm quy trình phát triển xuyên suốt.", indent_first=1)
add_table_caption(doc, "2.1")
add_simple_table(doc, ["Giai đoạn RUP", "Mục tiêu chính theo [1]", "Hoạt động áp dụng trong đề tài"], [
    ("Inception\n(Khởi tạo)", "Xác lập phạm vi và tầm nhìn dự án.",
     "Khảo sát nghiệp vụ bán lẻ điện thoại trực tuyến thực tế, xác định mục tiêu, phạm vi đề tài (Chương 1)."),
    ("Elaboration\n(Phân tích/Thiết kế)", "Phân tích bài toán, loại bỏ các rủi ro lớn, xây dựng Use Case và Mô hình khái niệm.",
     "Xác định actor, xây dựng Use Case Diagram, đặc tả Use Case và Class Diagram khái niệm (Chương 3, Chương 4)."),
    ("Construction\n(Xây dựng)", "Hoàn thiện thiết kế chi tiết và cài đặt phần lớn hệ thống.",
     "Thiết kế Sequence/Activity/State/Package/Component/Deployment Diagram, thiết kế CSDL và cài đặt mã nguồn (Chương 4, Chương 5)."),
    ("Transition\n(Chuyển giao)", "Kiểm thử, hoàn thiện và đưa hệ thống vào sử dụng.",
     "Kiểm thử API bằng Postman, kiểm thử giao diện trên dữ liệu thật, hoàn thiện báo cáo (Chương 5, Chương 6)."),
], widths=[3.5, 6, 6.7])

add_heading(doc, "2.4. Khung đặc tả Use Case", level=2)
add_para(doc,
    "Theo [1], UML không quy định cụ thể cấu trúc của tài liệu đặc tả (description) cho mỗi Use "
    "Case, việc này tùy thuộc vào quy ước của từng dự án/công ty. Khóa luận áp dụng khung đặc tả "
    "Use Case được đề xuất trong tài liệu [1], gồm các trường: Tên Use Case, Mô tả ngắn, Điều kiện "
    "trước (Pre-condition), Điều kiện sau (Post-condition), Luồng chính (Main Flow), Luồng phụ "
    "(Alternate Flow) và Luồng ngoại lệ (Exception Flow). Khung đặc tả này được sử dụng để mô tả "
    "chi tiết các Use Case quan trọng tại Chương 3.", indent_first=1)

add_heading(doc, "2.5. Nền tảng công nghệ liên quan", level=2)
add_para(doc,
    "Ngoài phương pháp luận OOAD/UML, hệ thống PhoneStore được xây dựng dựa trên một số nền tảng "
    "công nghệ web và trí tuệ nhân tạo hiện đại, được trình bày ngắn gọn dưới đây; chi tiết công "
    "nghệ và phiên bản cụ thể đã sử dụng được trình bày tại Chương 5.", indent_first=1)
add_para(doc, "REST API và WebSocket", bold=True, space_after=2)
add_para(doc,
    "REST (Representational State Transfer) là kiến trúc thiết kế API phổ biến dựa trên giao thức "
    "HTTP, trong đó mỗi tài nguyên (resource) được truy cập qua một địa chỉ URL và các phương thức "
    "chuẩn (GET, POST, PUT, DELETE) [2][4]. Bên cạnh REST, hệ thống còn sử dụng WebSocket (qua thư "
    "viện Socket.IO [6]) để hỗ trợ giao tiếp hai chiều theo thời gian thực, phục vụ chức năng chat "
    "trực tiếp và thông báo.", indent_first=1)
add_para(doc, "JWT và xác thực người dùng", bold=True, space_after=2)
add_para(doc,
    "JSON Web Token (JWT) là một chuẩn mở để truyền tải thông tin xác thực giữa các bên dưới dạng "
    "một đối tượng JSON đã được ký số, cho phép máy chủ xác minh người dùng mà không cần lưu trạng "
    "thái phiên đăng nhập (stateless). Hệ thống sử dụng cơ chế cặp access token (thời hạn ngắn) và "
    "refresh token (thời hạn dài) để vừa đảm bảo an toàn vừa duy trì trải nghiệm đăng nhập liên tục.", indent_first=1)
add_para(doc, "Cơ sở dữ liệu NoSQL (MongoDB)", bold=True, space_after=2)
add_para(doc,
    "MongoDB [4] là hệ quản trị cơ sở dữ liệu NoSQL hướng văn bản (document-oriented), lưu dữ liệu "
    "dưới dạng các văn bản BSON linh hoạt thay vì các bảng quan hệ cố định. Mô hình này phù hợp với "
    "các thực thể có cấu trúc lồng nhau và thường thay đổi như sản phẩm có nhiều biến thể hoặc đơn "
    "hàng có nhiều mục hàng, vốn rất phổ biến trong hệ thống PhoneStore.", indent_first=1)
add_para(doc, "Mô hình sinh văn bản có truy xuất dữ liệu thật (RAG)", bold=True, space_after=2)
add_para(doc,
    "Retrieval-Augmented Generation (RAG) là kỹ thuật kết hợp một mô hình ngôn ngữ lớn (LLM) với "
    "một bước truy xuất dữ liệu thật trước khi sinh câu trả lời, nhằm giảm thiểu hiện tượng mô hình "
    "\"bịa\" thông tin (hallucination). Trong hệ thống PhoneStore, trước khi LLM (chạy qua Ollama) "
    "sinh câu trả lời, dịch vụ chatbot sẽ gọi API thật của hệ thống để lấy dữ liệu sản phẩm, giá, "
    "tồn kho hiện tại và đưa vào ngữ cảnh (context) của câu hỏi, đảm bảo câu trả lời luôn phản ánh "
    "đúng dữ liệu kinh doanh thực tế tại thời điểm hỏi.", indent_first=1)

# ===================== CHUONG 3 =====================
add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG", level=1, page_break_before=True)

add_heading(doc, "3.1. Xác định Actor", level=2)
add_para(doc,
    "Qua khảo sát nghiệp vụ bán lẻ điện thoại trực tuyến, hệ thống PhoneStore xác định ba actor "
    "chính, trong đó Customer kế thừa toàn bộ quyền của Guest, và Admin kế thừa toàn bộ quyền của "
    "Customer (quan hệ generalization), thể hiện việc một quản trị viên vẫn có thể thực hiện được "
    "các chức năng mua hàng thông thường.", indent_first=1)
add_table_caption(doc, "3.1")
add_simple_table(doc, ["Actor", "Mô tả"], [
    ("Khách (Guest)", "Người dùng chưa đăng nhập: xem sản phẩm, tìm kiếm, so sánh, đăng ký/đăng nhập, chat hỏi đáp với chatbot."),
    ("Khách hàng (Customer)", "Người dùng đã đăng nhập với vai trò mua hàng: quản lý giỏ hàng, đặt hàng, theo dõi đơn hàng, đổi trả, đánh giá, wishlist, điểm thưởng, ví điện tử."),
    ("Quản trị viên (Admin/Staff)", "Người vận hành cửa hàng: quản lý sản phẩm, danh mục, đơn hàng, người dùng, khuyến mãi, đánh giá, đổi trả và hỗ trợ chat trực tiếp."),
    ("Chatbot AI", "Hệ thống tự động hỗ trợ tư vấn, đóng vai trò actor phụ trong Use Case \"Chat với hệ thống\"."),
    ("Cổng thanh toán VNPay", "Hệ thống bên thứ ba xử lý thanh toán trực tuyến, đóng vai trò actor phụ trong Use Case \"Thanh toán VNPay\"."),
], widths=[4.5, 11.2])

add_heading(doc, "3.2. Use Case Diagram tổng thể", level=2)
add_para(doc,
    "Hình 3.1 trình bày Use Case Diagram tổng thể của hệ thống PhoneStore, được nhóm theo 7 nhóm "
    "chức năng: Tìm kiếm & xem sản phẩm, Tài khoản, Giỏ hàng & Thanh toán, Đơn hàng & Đổi trả, "
    "Tương tác khách hàng, Điểm thưởng & Ví, và Quản trị hệ thống.", indent_first=1)
add_figure(doc, "3.1")
add_para(doc,
    "Trong sơ đồ, quan hệ <<include>> giữa \"Đặt hàng (checkout)\" và \"Thanh toán VNPay\" thể hiện "
    "việc thanh toán VNPay là một bước con bắt buộc xảy ra khi khách hàng chọn phương thức thanh "
    "toán này. Tương tự, \"Chat với hệ thống\" <<include>> sự tham gia của actor Chatbot AI, vì mọi "
    "phiên chat đều được Bot xử lý trước; khi cần hỗ trợ chuyên sâu hơn, Use Case \"Hỗ trợ chat trực "
    "tiếp\" của Admin sẽ <<extend>> (mở rộng) phiên chat đó bằng sự tham gia trực tiếp của nhân viên.", indent_first=1)

add_heading(doc, "3.3. Tổng hợp Use Case theo nhóm chức năng", level=2)
add_table_caption(doc, "3.2")
add_simple_table(doc, ["STT", "Use Case", "Actor chính"], [
    (1, "Xem danh sách sản phẩm", "Guest"), (2, "Tìm kiếm / lọc sản phẩm", "Guest"),
    (3, "Xem chi tiết sản phẩm", "Guest"), (4, "So sánh sản phẩm", "Guest"),
    (5, "Xem Flash Sale", "Guest"), (6, "Đăng ký", "Guest"),
    (7, "Đăng nhập (local / Google)", "Guest"), (8, "Quên mật khẩu", "Guest"),
    (9, "Quản lý hồ sơ & địa chỉ", "Customer"), (10, "Quản lý giỏ hàng", "Customer"),
    (11, "Áp mã giảm giá", "Customer"), (12, "Đặt hàng (checkout)", "Customer"),
    (13, "Thanh toán VNPay", "Customer"), (14, "Xem / hủy đơn hàng", "Customer"),
    (15, "Yêu cầu trả hàng", "Customer"), (16, "Đánh giá sản phẩm", "Customer"),
    (17, "Quản lý wishlist", "Customer"), (18, "Chat với hệ thống", "Guest/Customer"),
    (19, "Xem điểm thưởng & hạng thành viên", "Customer"), (20, "Đổi điểm sang mã giảm giá", "Customer"),
    (21, "Quản lý ví điện tử (nạp / rút)", "Customer"), (22, "Quản lý sản phẩm & danh mục", "Admin"),
    (23, "Quản lý đơn hàng", "Admin"), (24, "Quản lý người dùng", "Admin"),
    (25, "Quản lý đổi trả", "Admin"), (26, "Quản lý mã giảm giá & Flash Sale", "Admin"),
    (27, "Quản lý đánh giá", "Admin"), (28, "Xem báo cáo thống kê", "Admin"),
], widths=[1.3, 9.5, 4.9])

add_heading(doc, "3.4. Đặc tả các Use Case chính", level=2)
add_para(doc,
    "Phần này trình bày đặc tả chi tiết cho 8 Use Case tiêu biểu, đại diện cho các nhóm chức năng "
    "quan trọng nhất của hệ thống (tài khoản, mua hàng, khuyến mãi, điểm thưởng, đổi trả, chatbot "
    "và quản trị), theo khung đặc tả đã trình bày tại mục 2.4.", indent_first=1)

add_usecase_table(doc, "3.3", "Đăng ký / Đăng nhập tài khoản",
    "Cho phép Guest tạo tài khoản mới hoặc đăng nhập vào hệ thống bằng email/mật khẩu hoặc tài khoản Google.",
    "Guest",
    "Người dùng chưa đăng nhập; đối với đăng ký, email chưa từng được sử dụng trong hệ thống.",
    "Người dùng có phiên đăng nhập hợp lệ (access token, refresh token) và được điều hướng vào hệ thống.",
    ["1. Người dùng chọn Đăng nhập, nhập email và mật khẩu.",
     "2. Hệ thống kiểm tra thông tin trong cơ sở dữ liệu.",
     "3. Nếu hợp lệ, hệ thống sinh JWT access token và refresh token, trả về thông tin người dùng.",
     "4. Hệ thống lưu phiên đăng nhập và chuyển hướng vào trang chủ."],
    ["Đăng ký mới: người dùng nhập họ tên, email, mật khẩu; hệ thống gửi mã OTP xác thực email "
     "trước khi tạo tài khoản hoàn chỉnh.",
     "Đăng nhập bằng Google: người dùng xác thực qua Google OAuth, hệ thống tự tạo/liên kết tài "
     "khoản tương ứng."],
    ["Sai email/mật khẩu: hệ thống hiển thị thông báo lỗi, giới hạn số lần thử trong 15 phút "
     "(chống brute-force).",
     "OTP hết hạn hoặc sai: hệ thống yêu cầu gửi lại mã mới."])

add_usecase_table(doc, "3.4", "Đặt hàng (Checkout)",
    "Cho phép khách hàng tạo đơn hàng từ giỏ hàng hiện tại, có thể áp dụng mã giảm giá, Flash Sale "
    "và điểm thưởng trước khi xác nhận thanh toán.",
    "Customer",
    "Khách hàng đã đăng nhập, giỏ hàng có ít nhất một sản phẩm còn tồn kho.",
    "Đơn hàng được tạo với trạng thái \"pending\", tồn kho được trừ tương ứng, điểm/coupon (nếu có) "
    "được ghi nhận sử dụng.",
    ["1. Khách hàng chọn sản phẩm trong giỏ hàng, nhập/chọn địa chỉ giao hàng.",
     "2. Khách hàng chọn phương thức vận chuyển và phương thức thanh toán.",
     "3. Hệ thống kiểm tra tồn kho, tính lại tổng tiền (đã gồm giảm giá, điểm sử dụng, phí vận "
     "chuyển).",
     "4. Khách hàng xác nhận đặt hàng; hệ thống tạo Order và trừ tồn kho.",
     "5. Hệ thống hiển thị trang xác nhận đặt hàng thành công."],
    ["Áp dụng mã giảm giá hợp lệ trước khi đặt hàng.",
     "Sử dụng một phần điểm thưởng để giảm trừ vào tổng tiền.",
     "Chọn thanh toán bằng ví điện tử: hệ thống trừ trực tiếp walletBalance khi tạo đơn."],
    ["Hết hàng giữa lúc đặt: hệ thống hủy thao tác và thông báo hết hàng cho khách.",
     "Mã giảm giá hết lượt sử dụng hoặc không hợp lệ ở thời điểm đặt: hệ thống từ chối áp dụng.",
     "Lỗi trong quá trình tạo đơn: hệ thống hoàn tác (rollback) các thay đổi atomic đã thực hiện "
     "trước đó (tồn kho, lượt dùng mã giảm giá...)."])

add_usecase_table(doc, "3.5", "Áp dụng mã giảm giá",
    "Cho phép khách hàng nhập mã giảm giá (coupon) trong quá trình thanh toán để được giảm trừ vào "
    "tổng tiền đơn hàng.",
    "Customer",
    "Khách hàng đang ở trang giỏ hàng hoặc thanh toán; tồn tại mã giảm giá hợp lệ.",
    "Tổng tiền đơn hàng được giảm trừ theo giá trị/loại của mã; trạng thái áp dụng được hiển thị.",
    ["1. Khách hàng nhập mã giảm giá.",
     "2. Hệ thống kiểm tra thời hạn, giá trị đơn hàng tối thiểu, hạng thành viên và số lượt còn lại "
     "của mã.",
     "3. Nếu hợp lệ, hệ thống tính số tiền được giảm và cập nhật lại tổng tiền."],
    ["Mã giảm giá riêng cho một khách hàng cụ thể (allowedUserId) hoặc theo nhóm hạng thành viên."],
    ["Mã không tồn tại, đã hết hạn, đã hết lượt sử dụng hoặc đơn hàng chưa đạt giá trị tối thiểu: "
     "hệ thống hiển thị thông báo lỗi tương ứng và không áp dụng giảm giá."])

add_usecase_table(doc, "3.6", "Đổi điểm thưởng sang mã giảm giá",
    "Cho phép khách hàng dùng điểm thưởng tích lũy để đổi sang một mã giảm giá sử dụng cho lần mua "
    "tiếp theo.",
    "Customer",
    "Khách hàng có số điểm thưởng (loyaltyPoints) lớn hơn hoặc bằng mức tối thiểu cho phép đổi (50 "
    "điểm).",
    "Điểm thưởng của khách hàng giảm tương ứng; một mã giảm giá mới (loại fixed) được tạo và gán "
    "riêng cho khách hàng đó.",
    ["1. Khách hàng nhập số điểm muốn đổi.",
     "2. Hệ thống kiểm tra và trừ điểm thưởng (thao tác atomic, chỉ thành công nếu đủ điểm).",
     "3. Hệ thống sinh mã giảm giá ngẫu nhiên, gán allowedUserId cho khách hàng, giới hạn 1 lượt sử "
     "dụng, hết hạn sau 30 ngày.",
     "4. Hệ thống trả về mã giảm giá mới và số điểm còn lại."],
    [],
    ["Không đủ điểm để đổi: hệ thống từ chối và hiển thị thông báo lỗi."])

add_usecase_table(doc, "3.7", "Yêu cầu trả hàng",
    "Cho phép khách hàng gửi yêu cầu trả lại một phần hoặc toàn bộ sản phẩm trong đơn hàng đã giao, "
    "kèm lý do trả hàng.",
    "Customer (gửi yêu cầu), Admin (duyệt/xử lý)",
    "Đơn hàng đã ở trạng thái \"delivered\" và còn trong thời hạn 7 ngày kể từ ngày giao hàng; chưa "
    "có yêu cầu trả hàng nào khác đang xử lý cho đơn này.",
    "Yêu cầu trả hàng được tạo (trạng thái pending); sau khi Admin duyệt, đơn được hoàn tiền (vào ví "
    "hoặc chuyển khoản) và điểm thưởng/hạng thành viên được tính lại tương ứng.",
    ["1. Khách hàng chọn đơn hàng đã giao, chọn sản phẩm cần trả và nhập lý do.",
     "2. Hệ thống tạo ReturnRequest và chuyển trạng thái đơn hàng thành \"return_requested\".",
     "3. Admin xem xét, duyệt yêu cầu và chọn phương thức hoàn tiền.",
     "4. Hệ thống trừ điểm thưởng theo tỉ lệ giá trị hoàn trả, cập nhật lại hạng thành viên (có thể "
     "bị hạ hạng), ghi StockLog nhập lại tồn kho và chuyển trạng thái đơn hàng thành \"returned\"."],
    ["Hoàn tiền vào ví điện tử của khách hàng (ghi nhận WalletTransaction).",
     "Hoàn tiền qua chuyển khoản ngân hàng (lưu thông tin tham chiếu hoàn tiền)."],
    ["Admin từ chối yêu cầu: đơn hàng trở lại trạng thái \"delivered\", không hoàn tiền.",
     "Quá thời hạn 7 ngày hoặc đơn chưa giao: hệ thống không cho phép tạo yêu cầu trả hàng."])

add_usecase_table(doc, "3.8", "Chat với hệ thống (Chatbot AI)",
    "Cho phép khách hàng đặt câu hỏi về sản phẩm, giá, tồn kho... và nhận câu trả lời tự động từ "
    "chatbot AI dựa trên dữ liệu thật của hệ thống.",
    "Guest/Customer, Chatbot AI",
    "Người dùng mở khung chat trên giao diện web.",
    "Câu trả lời được hiển thị cho khách hàng; lịch sử hội thoại được lưu vào ChatSession/"
    "ChatMessage.",
    ["1. Khách hàng gửi câu hỏi qua khung chat (kết nối Socket.IO).",
     "2. Hệ thống Node.js chuyển câu hỏi tới dịch vụ chatbot Python.",
     "3. Dịch vụ chatbot phân tích ý định câu hỏi, truy vấn dữ liệu thật tương ứng (sản phẩm, giá, "
     "tồn kho...) qua API của hệ thống Node.js.",
     "4. Dịch vụ chatbot ghép dữ liệu thật vào ngữ cảnh và gửi cho mô hình ngôn ngữ lớn (LLM) để "
     "sinh câu trả lời tự nhiên.",
     "5. Câu trả lời được trả về và hiển thị cho khách hàng, kèm các nút điều hướng (nếu có)."],
    ["Khi cần hỗ trợ chuyên sâu, nhân viên có thể tiếp nhận trực tiếp phiên chat (Use Case \"Hỗ trợ "
     "chat trực tiếp\")."],
    ["Dịch vụ AI (Ollama) không phản hồi hoặc lỗi: hệ thống hiển thị thông báo lỗi tạm thời cho "
     "khách hàng."])

add_usecase_table(doc, "3.9", "Quản lý sản phẩm (Admin)",
    "Cho phép quản trị viên thêm, sửa, xóa sản phẩm cùng các biến thể (màu sắc, dung lượng, giá, "
    "tồn kho) và danh mục/thương hiệu liên quan.",
    "Admin/Staff (có quyền manage_products)",
    "Người dùng đã đăng nhập với vai trò admin hoặc staff có quyền quản lý sản phẩm.",
    "Dữ liệu sản phẩm/biến thể được cập nhật trong cơ sở dữ liệu và phản ánh ngay trên giao diện "
    "khách hàng.",
    ["1. Quản trị viên truy cập trang Quản lý sản phẩm.",
     "2. Quản trị viên thêm sản phẩm mới hoặc chọn sản phẩm cần sửa.",
     "3. Quản trị viên nhập/cập nhật thông tin sản phẩm và các biến thể (giá, tồn kho, ảnh).",
     "4. Hệ thống lưu thay đổi và cập nhật lại trạng thái hiển thị sản phẩm."],
    ["Cập nhật tồn kho hàng loạt; ghi nhận StockLog cho mỗi thay đổi tồn kho thủ công."],
    ["Xóa sản phẩm đang có trong đơn hàng/giỏ hàng của khách: hệ thống chỉ cho phép ẩn (status) "
     "thay vì xóa cứng để bảo toàn dữ liệu lịch sử đơn hàng."])

add_usecase_table(doc, "3.10", "Quản lý đơn hàng (Admin)",
    "Cho phép quản trị viên xem danh sách đơn hàng, cập nhật trạng thái xử lý đơn (xác nhận, chuẩn "
    "bị hàng, giao hàng, hoàn tất) và xem chi tiết từng đơn.",
    "Admin/Staff (có quyền manage_orders)",
    "Người dùng đã đăng nhập với vai trò admin hoặc staff có quyền quản lý đơn hàng.",
    "Trạng thái đơn hàng được cập nhật; nếu chuyển sang \"delivered\", điểm thưởng được tích lũy "
    "cho khách hàng và hạng thành viên được tính lại.",
    ["1. Quản trị viên xem danh sách đơn hàng, có thể lọc theo trạng thái.",
     "2. Quản trị viên chọn một đơn hàng và cập nhật trạng thái xử lý kế tiếp.",
     "3. Hệ thống kiểm tra tính hợp lệ của bước chuyển trạng thái theo vòng đời đơn hàng (xem Hình "
     "4.12).",
     "4. Khi chuyển trạng thái thành \"delivered\", hệ thống tự động tính và cộng điểm thưởng, cập "
     "nhật hạng thành viên, gửi thông báo cho khách hàng."],
    ["Hủy đơn hàng theo yêu cầu của khách (trước khi giao) hoặc tự động hủy nếu quá hạn chờ thanh "
     "toán."],
    ["Cố gắng chuyển sang trạng thái không hợp lệ theo vòng đời đơn hàng (ví dụ từ \"pending\" sang "
     "\"delivered\" trực tiếp): hệ thống từ chối thao tác."])

add_heading(doc, "3.5. Yêu cầu phi chức năng", level=2)
add_para(doc, "Hiệu năng", bold=True, space_after=2)
add_para(doc,
    "Hệ thống áp dụng giới hạn tốc độ truy cập API (rate limiting) ở ba mức: tối đa 100 request/"
    "phút cho API thông thường, 10 lần đăng nhập/15 phút và 5 yêu cầu OTP/giờ cho mỗi địa chỉ IP, "
    "nhằm đảm bảo thời gian phản hồi ổn định và hạn chế quá tải máy chủ.", indent_first=1)
add_para(doc, "An toàn và bảo mật", bold=True, space_after=2)
add_para(doc,
    "Mật khẩu người dùng được mã hóa một chiều bằng bcrypt trước khi lưu trữ; phiên đăng nhập sử "
    "dụng cặp JWT access token/refresh token; các thao tác nhạy cảm (đăng ký, quên mật khẩu) yêu "
    "cầu xác thực OTP gửi qua email; phân quyền truy cập theo vai trò (role-based access control) "
    "được áp dụng chặt chẽ cho khu vực quản trị; các tiêu đề bảo mật HTTP được thiết lập qua "
    "middleware Helmet và chính sách CORS giới hạn nguồn gốc truy cập hợp lệ.", indent_first=1)
add_para(doc, "Tính khả dụng", bold=True, space_after=2)
add_para(doc,
    "Các tác vụ định kỳ (cron job) tự động hủy đơn hàng quá hạn chờ thanh toán và giải phóng tồn "
    "kho đã giữ chỗ, tránh tình trạng tồn kho bị \"khóa\" vô thời hạn do khách hàng bỏ ngang quá "
    "trình thanh toán.", indent_first=1)
add_para(doc, "Khả năng mở rộng và bảo trì", bold=True, space_after=2)
add_para(doc,
    "Kiến trúc hệ thống tách dịch vụ chatbot AI thành một tiến trình Python độc lập, giao tiếp với "
    "hệ thống chính qua REST API, cho phép nâng cấp hoặc thay thế mô hình AI mà không ảnh hưởng đến "
    "hệ thống lõi. Mã nguồn được tổ chức theo từng module rõ ràng (routes/controllers/models) như "
    "trình bày trong Package Diagram (Hình 4.13), giúp dễ bảo trì và mở rộng.", indent_first=1)

# ===================== CHUONG 4 =====================
add_heading(doc, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", level=1, page_break_before=True)

add_heading(doc, "4.1. Kiến trúc tổng thể", level=2)
add_para(doc,
    "Hệ thống PhoneStore được thiết kế theo kiến trúc nhiều tầng (xem lại Hình 1.1): tầng giao "
    "diện là một ứng dụng React một trang (Single Page Application), giao tiếp với tầng ứng dụng "
    "là máy chủ Node.js/Express thông qua REST API (cho các thao tác CRUD thông thường) và "
    "WebSocket/Socket.IO (cho chat và thông báo thời gian thực). Tầng ứng dụng còn bao gồm một "
    "dịch vụ chatbot AI viết bằng Python/FastAPI, giao tiếp với máy chủ Node.js qua HTTP để lấy dữ "
    "liệu thật, và với công cụ Ollama để sinh câu trả lời bằng mô hình ngôn ngữ lớn đã fine-tune "
    "riêng cho lĩnh vực bán điện thoại. Toàn bộ dữ liệu nghiệp vụ được lưu trong MongoDB thông qua "
    "thư viện Mongoose. Ngoài ra, hệ thống còn có các tác vụ định kỳ (cron job) chạy độc lập để tự "
    "động hủy đơn hàng quá hạn, giải phóng tồn kho giữ chỗ và sinh mã giảm giá sinh nhật cho khách "
    "hàng.", indent_first=1)

add_heading(doc, "4.2. Mô hình hóa khái niệm - Class Diagram tổng quan", level=2)
add_para(doc,
    "Hình 4.1 trình bày Class Diagram tổng quan của toàn hệ thống, gồm 23 lớp thực thể tương ứng "
    "với 23 collection dữ liệu trong MongoDB (chi tiết tại mục 4.12), được nhóm thành ba miền "
    "nghiệp vụ chính: (1) Sản phẩm và Mua hàng, (2) Khuyến mãi và Tài chính, (3) Hỗ trợ và Tương "
    "tác khách hàng. Lớp User đóng vai trò trung tâm, có quan hệ với hầu hết các thực thể còn lại "
    "vì hầu hết nghiệp vụ của hệ thống đều xuất phát từ một người dùng cụ thể.", indent_first=1)
add_figure(doc, "4.1")

add_heading(doc, "4.3. Class Diagram - Nhóm Sản phẩm và Mua hàng", level=2)
add_para(doc,
    "Hình 4.2 mô tả chi tiết các lớp thuộc miền Sản phẩm và Mua hàng. Một Product có thể có nhiều "
    "ProductVariant (biến thể theo màu sắc, dung lượng), mỗi biến thể có giá, giá khuyến mãi và "
    "tồn kho riêng. Cart và Order đều chứa danh sách các mục hàng dưới dạng lớp nhúng (embedded "
    "class, đánh dấu <<embedded>>) là CartItem và OrderItem — đây là cách mô hình hóa phù hợp với "
    "MongoDB, nơi dữ liệu thường được lồng trực tiếp vào văn bản cha thay vì tách bảng như trong cơ "
    "sở dữ liệu quan hệ. Order lưu lại snapshot giá tại thời điểm đặt hàng (price trong OrderItem) "
    "để đảm bảo lịch sử đơn hàng không bị ảnh hưởng khi giá sản phẩm thay đổi sau đó.", indent_first=1)
add_figure(doc, "4.2")

add_heading(doc, "4.4. Class Diagram - Nhóm Khuyến mãi và Tài chính", level=2)
add_para(doc,
    "Hình 4.3 mô tả các lớp liên quan đến khuyến mãi (FlashSale, Coupon) và tài chính (Payment, "
    "WalletTransaction, TopupRequest, WithdrawalRequest, StockLog). FlashSale có thể áp dụng theo "
    "biến thể sản phẩm cụ thể hoặc theo cả một danh mục; số lượng đã bán (sold) được cập nhật bằng "
    "thao tác atomic có điều kiện ($expr) để tránh bán vượt số lượng giới hạn khi nhiều khách hàng "
    "đặt hàng đồng thời. Coupon hỗ trợ nhiều điều kiện áp dụng: giá trị đơn tối thiểu, giới hạn "
    "lượt dùng toàn hệ thống/từng người dùng, giới hạn theo hạng thành viên hoặc theo một người "
    "dùng cụ thể (allowedUserId) — đây chính là cơ chế dùng để cấp mã giảm giá khi khách hàng đổi "
    "điểm thưởng. StockLog ghi lại toàn bộ lịch sử biến động tồn kho (bán, trả hàng, admin điều "
    "chỉnh) để phục vụ truy vết.", indent_first=1)
add_figure(doc, "4.3")

add_heading(doc, "4.5. Class Diagram - Nhóm Hỗ trợ và Tương tác khách hàng", level=2)
add_para(doc,
    "Hình 4.4 mô tả các lớp phục vụ tương tác khách hàng: Review (đánh giá sản phẩm, có thể gắn "
    "với một đơn hàng cụ thể để xác minh đã mua hàng), Wishlist, ReturnRequest (yêu cầu trả hàng), "
    "ChatSession/ChatMessage (phiên chat và từng tin nhắn, hỗ trợ cả tin nhắn từ Bot và từ nhân "
    "viên), Notification, OTP và các lớp nội dung tĩnh Banner, ServiceBadge phục vụ hiển thị trang "
    "chủ.", indent_first=1)
add_figure(doc, "4.4")

add_heading(doc, "4.6. Thiết kế tương tác đối tượng", level=2)
add_para(doc,
    "Phần này trình bày thiết kế chi tiết các luồng tương tác đối tượng quan trọng nhất của hệ "
    "thống bằng Sequence Diagram (và một Collaboration Diagram minh họa cho luồng đặt hàng), theo "
    "đúng định nghĩa tại mục 2.2.3 và 2.2.4.", indent_first=1)

add_heading(doc, "4.6.1. Đặt hàng (Checkout)", level=3)
add_para(doc,
    "Hình 4.5 mô tả chi tiết luồng xử lý khi khách hàng đặt hàng. Sau khi kiểm tra tồn kho, hệ "
    "thống tuần tự kiểm tra và áp dụng (nếu có) Flash Sale, mã giảm giá và điểm thưởng — mỗi bước "
    "đều là một thao tác atomic trên MongoDB để tránh xung đột khi có nhiều yêu cầu đồng thời. Nếu "
    "khách hàng chọn thanh toán bằng ví điện tử, số dư được trừ ngay khi tạo đơn. Đơn hàng được tạo "
    "với snapshot giá và thông tin giảm giá tại thời điểm đặt, sau đó hệ thống trừ tồn kho cho từng "
    "sản phẩm và ghi StockLog tương ứng. Nếu một bước bất kỳ thất bại giữa chừng (ví dụ hết hàng "
    "ngay trước khi trừ kho), hệ thống sẽ hoàn tác các thao tác atomic đã thực hiện trước đó để "
    "đảm bảo tính nhất quán dữ liệu.", indent_first=1)
add_figure(doc, "4.5")
add_para(doc,
    "Hình 4.6 trình bày lại cùng luồng nghiệp vụ trên dưới dạng Collaboration Diagram, nhấn mạnh "
    "vào các đối tượng tham gia cộng tác và các thông điệp được đánh số thứ tự giữa chúng, là một "
    "góc nhìn khác bổ sung cho Sequence Diagram theo đúng quan hệ giữa hai loại sơ đồ này như trình "
    "bày tại mục 2.2.3 và 2.2.4.", indent_first=1)
add_figure(doc, "4.6")

add_heading(doc, "4.6.2. Giao hàng thành công và tích điểm thưởng", level=3)
add_para(doc,
    "Hình 4.7 mô tả luồng xử lý khi quản trị viên cập nhật trạng thái đơn hàng thành \"delivered\". "
    "Hệ thống tính số điểm thưởng được cộng theo công thức floor(totalPrice / 1000) nhân với hệ số "
    "theo hạng thành viên hiện tại của khách hàng (bronze = 1, silver = 1.1, gold = 1.2, platinum = "
    "1.5), sau đó cộng điểm và tổng chi tiêu (totalSpent) cho khách hàng. Nếu tổng chi tiêu vượt "
    "ngưỡng của một hạng cao hơn (5 triệu/20 triệu/50 triệu đồng), hạng thành viên của khách hàng "
    "được nâng lên tương ứng và một thông báo được gửi tới khách hàng.", indent_first=1)
add_figure(doc, "4.7")

add_heading(doc, "4.6.3. Yêu cầu trả hàng và hoàn tiền", level=3)
add_para(doc,
    "Hình 4.8 mô tả hai giai đoạn của luồng trả hàng: giai đoạn khách hàng gửi yêu cầu (kiểm tra "
    "đơn đã giao trong vòng 7 ngày) và giai đoạn quản trị viên xử lý (duyệt yêu cầu, tính lại điểm "
    "thưởng/hạng thành viên theo tỉ lệ giá trị hoàn trả, hoàn tiền vào ví hoặc qua chuyển khoản, "
    "ghi nhận lại tồn kho). Việc tính lại điểm thưởng theo đúng tỉ lệ giá trị hàng được trả (thay "
    "vì trừ toàn bộ điểm của đơn hàng) đảm bảo tính công bằng khi khách hàng chỉ trả một phần đơn "
    "hàng.", indent_first=1)
add_figure(doc, "4.8")

add_heading(doc, "4.6.4. Đổi điểm thưởng sang mã giảm giá", level=3)
add_para(doc,
    "Hình 4.9 mô tả luồng đổi điểm thưởng: hệ thống trừ điểm bằng một thao tác atomic (chỉ thành "
    "công nếu khách hàng có đủ điểm), sau đó sinh một mã giảm giá loại cố định (fixed) với giá trị "
    "tương ứng số điểm đã đổi (mỗi điểm tương đương 1.000 đồng), gán riêng cho khách hàng đó, giới "
    "hạn một lượt sử dụng và có hiệu lực trong 30 ngày.", indent_first=1)
add_figure(doc, "4.9")

add_heading(doc, "4.6.5. Chatbot AI trả lời câu hỏi khách hàng (RAG)", level=3)
add_para(doc,
    "Hình 4.10 mô tả luồng xử lý một câu hỏi của khách hàng theo cơ chế RAG đã trình bày ở mục 2.5. "
    "Câu hỏi được gửi từ giao diện qua Socket.IO tới Node.js, rồi chuyển tiếp cho dịch vụ Python. "
    "Dịch vụ này phân tích ý định câu hỏi (intent.py), gọi lại API thật của hệ thống Node.js để lấy "
    "dữ liệu sản phẩm/giá/tồn kho tương ứng, ghép dữ liệu đó vào phần \"ngữ cảnh tra cứu\" của "
    "system prompt, rồi mới gửi cho mô hình ngôn ngữ lớn (LLM, chạy qua Ollama) để sinh câu trả "
    "lời. Nguyên tắc thiết kế quan trọng ở đây là LLM chỉ được sử dụng dữ liệu trong phần ngữ cảnh "
    "tra cứu để trả lời các câu hỏi về giá/tồn kho/màu sắc cụ thể, không được tự suy diễn từ kiến "
    "thức huấn luyện trước đó — đây chính là biện pháp thiết kế nhằm hạn chế hiện tượng \"bịa\" "
    "thông tin của LLM.", indent_first=1)
add_figure(doc, "4.10")

add_heading(doc, "4.7. Activity Diagram - Luồng đặt hàng", level=2)
add_para(doc,
    "Hình 4.11 trình bày Activity Diagram cho toàn bộ luồng đặt hàng dưới góc nhìn khách hàng, "
    "tổng hợp lại các điểm rẽ nhánh quan trọng đã phân tích ở Sequence Diagram (Hình 4.5) dưới một "
    "góc nhìn trực quan, dễ theo dõi hơn cho người không chuyên về kỹ thuật: kiểm tra mã giảm giá "
    "hợp lệ, kiểm tra đủ điểm thưởng, kiểm tra còn hàng, và xử lý kết quả thanh toán VNPay (nếu "
    "chọn phương thức này).", indent_first=1)
add_figure(doc, "4.11")

add_heading(doc, "4.8. State Diagram - Vòng đời đơn hàng", level=2)
add_para(doc,
    "Hình 4.12 mô tả các trạng thái hợp lệ và các phép chuyển trạng thái của đối tượng Order trong "
    "suốt vòng đời của nó, theo đúng nguyên tắc tại mục 2.2.6: một đơn hàng mới luôn bắt đầu ở "
    "trạng thái Pending, sau đó được Admin xác nhận (Confirmed), chuẩn bị hàng (Preparing), giao "
    "cho đơn vị vận chuyển (Shipping) và giao thành công (Delivered). Từ trạng thái Delivered, "
    "khách hàng có thể gửi yêu cầu trả hàng trong 7 ngày (ReturnRequested), dẫn tới trạng thái "
    "Returned nếu được duyệt hoặc trở lại Delivered nếu bị từ chối. Đơn hàng có thể bị hủy "
    "(Cancelled) ở các trạng thái Pending hoặc Confirmed. Việc kiểm soát chặt chẽ các phép chuyển "
    "này (ví dụ không cho phép chuyển trực tiếp từ Pending sang Delivered) giúp tránh các lỗi "
    "nghiệp vụ nghiêm trọng như đã đề cập tại mục 2.2.6.", indent_first=1)
add_figure(doc, "4.12")

add_heading(doc, "4.9. Package Diagram", level=2)
add_para(doc,
    "Hình 4.13 mô tả cách tổ chức module mã nguồn của ba thành phần Client, Server và Bot. Phía "
    "Client được chia thành pages, components, hooks, context và api; phía Server được chia thành "
    "routes, controllers, models, middlewares, jobs (cron) và utils/validations; phía Bot được "
    "chia thành main, intent, actions, templates và llm. Cách phân chia này được áp dụng trực tiếp "
    "khi tổ chức thư mục mã nguồn thật của hệ thống (xem mục 5.2).", indent_first=1)
add_figure(doc, "4.13")

add_heading(doc, "4.10. Component Diagram", level=2)
add_para(doc,
    "Hình 4.14 mô tả các thành phần logic (component) của hệ thống dưới dạng các dịch vụ nghiệp vụ "
    "(Auth, Product, Order, Payment, Loyalty, Notification) và dịch vụ Chatbot, mỗi dịch vụ giao "
    "tiếp với phần còn lại của hệ thống qua một giao diện (interface) tương ứng (IAuth, IProduct, "
    "IOrder...). Web Client chỉ phụ thuộc vào các interface này mà không phụ thuộc trực tiếp vào "
    "cách cài đặt bên trong từng dịch vụ — thể hiện nguyên tắc tách rời (decoupling) giữa giao diện "
    "và cài đặt.", indent_first=1)
add_figure(doc, "4.14")

add_heading(doc, "4.11. Deployment Diagram", level=2)
add_para(doc,
    "Hình 4.15 mô tả cách hệ thống được triển khai trên hạ tầng vật lý: máy khách (trình duyệt web "
    "chạy ứng dụng React đã build) giao tiếp với Application Server (chạy Node.js/Express qua HTTPS "
    "và WSS); Application Server giao tiếp với Database Server (MongoDB, qua giao thức Mongoose) "
    "và với AI Server (chạy Python/FastAPI và Ollama) qua REST nội bộ. Trong môi trường phát triển "
    "và kiểm thử của khóa luận, cả bốn node này được triển khai trên cùng một máy vật lý, nhưng "
    "kiến trúc tách rời theo node cho phép dễ dàng phân tách ra nhiều máy chủ vật lý/đám mây riêng "
    "biệt khi đưa vào môi trường sản xuất thực tế.", indent_first=1)
add_figure(doc, "4.15")

add_heading(doc, "4.12. Thiết kế cơ sở dữ liệu", level=2)
add_para(doc,
    "Hệ thống sử dụng MongoDB — một hệ quản trị cơ sở dữ liệu NoSQL hướng văn bản — thay vì cơ sở "
    "dữ liệu quan hệ, vì phần lớn dữ liệu nghiệp vụ (đơn hàng kèm danh sách sản phẩm, giỏ hàng, "
    "biến thể sản phẩm) có cấu trúc lồng nhau tự nhiên, phù hợp để nhúng trực tiếp (embedded "
    "document) thay vì phải JOIN nhiều bảng như trong mô hình quan hệ, đồng thời giúp việc đọc dữ "
    "liệu của một đơn hàng hoặc một sản phẩm chỉ cần một truy vấn duy nhất. Mỗi lớp trong Class "
    "Diagram (Hình 4.1) tương ứng với một collection dữ liệu trong MongoDB, được quản lý thông qua "
    "thư viện Mongoose ORM. Bảng 4.1 liệt kê toàn bộ 23 collection dữ liệu của hệ thống.", indent_first=1)
add_table_caption(doc, "4.1")
add_simple_table(doc, ["Collection", "Mô tả"], [
    ("User", "Tài khoản người dùng: thông tin cá nhân, vai trò, hạng thành viên, điểm thưởng, ví."),
    ("Brand", "Thương hiệu sản phẩm."),
    ("Category", "Danh mục sản phẩm, hỗ trợ danh mục cha-con."),
    ("Product", "Thông tin chung của sản phẩm (tên, mô tả, thông số, đánh giá)."),
    ("ProductVariant", "Biến thể sản phẩm theo màu sắc/dung lượng, giá và tồn kho riêng."),
    ("Cart", "Giỏ hàng hiện tại của từng người dùng."),
    ("Order", "Đơn hàng, kèm danh sách sản phẩm đã đặt (nhúng)."),
    ("FlashSale", "Chương trình giảm giá theo thời gian, theo biến thể hoặc danh mục."),
    ("Coupon", "Mã giảm giá và các điều kiện áp dụng."),
    ("Payment", "Giao dịch thanh toán của đơn hàng."),
    ("WalletTransaction", "Lịch sử biến động ví điện tử của người dùng."),
    ("TopupRequest", "Yêu cầu nạp tiền vào ví điện tử."),
    ("WithdrawalRequest", "Yêu cầu rút tiền từ ví điện tử."),
    ("StockLog", "Lịch sử biến động tồn kho (bán, trả hàng, điều chỉnh)."),
    ("Review", "Đánh giá sản phẩm của khách hàng."),
    ("Wishlist", "Danh sách sản phẩm yêu thích của khách hàng."),
    ("ReturnRequest", "Yêu cầu trả hàng/hoàn tiền."),
    ("ChatSession", "Phiên hội thoại giữa khách hàng và hệ thống hỗ trợ."),
    ("ChatMessage", "Từng tin nhắn trong một phiên hội thoại."),
    ("Notification", "Thông báo gửi tới người dùng."),
    ("OTP", "Mã xác thực dùng một lần (xác thực email, quên mật khẩu)."),
    ("Banner", "Banner quảng cáo hiển thị trên trang chủ."),
    ("ServiceBadge", "Huy hiệu dịch vụ/cam kết hiển thị trên trang chủ."),
], widths=[4.5, 11.2])

# ===================== CHUONG 5 =====================
add_heading(doc, "CHƯƠNG 5. CÀI ĐẶT VÀ KIỂM THỬ", level=1, page_break_before=True)

add_heading(doc, "5.1. Môi trường và công nghệ cài đặt", level=2)
add_para(doc,
    "Hệ thống được cài đặt và kiểm thử trên môi trường máy chủ cục bộ (local), sử dụng Node.js cho "
    "máy chủ ứng dụng chính, MongoDB cho cơ sở dữ liệu, và Python cho dịch vụ chatbot AI. Bảng 5.1, "
    "5.2 và 5.3 liệt kê chi tiết các thư viện/công nghệ thực tế đã sử dụng ở từng thành phần, được "
    "trích trực tiếp từ tệp khai báo phụ thuộc (package.json, requirements.txt) của dự án.", indent_first=1)

add_table_caption(doc, "5.1")
add_simple_table(doc, ["Thư viện", "Vai trò"], [
    ("React 18", "Thư viện xây dựng giao diện người dùng dạng component."),
    ("Vite 5", "Công cụ build và dev server cho ứng dụng React."),
    ("React Router DOM 6", "Định tuyến (routing) phía client."),
    ("Axios", "Gửi yêu cầu HTTP tới REST API của Server."),
    ("Socket.IO Client", "Giao tiếp thời gian thực (chat, thông báo) qua WebSocket."),
    ("TailwindCSS 3", "Xây dựng giao diện theo phương pháp utility-first CSS."),
    ("React Hook Form", "Quản lý và validate form phía client."),
    ("Recharts", "Vẽ biểu đồ thống kê doanh thu trong trang quản trị."),
    ("@react-oauth/google", "Đăng nhập bằng tài khoản Google (OAuth2)."),
    ("i18next / react-i18next", "Hỗ trợ đa ngôn ngữ cho giao diện."),
    ("lucide-react", "Bộ icon dùng trong giao diện."),
], widths=[5, 10.7])

add_table_caption(doc, "5.2")
add_simple_table(doc, ["Thư viện", "Vai trò"], [
    ("Express 4", "Framework xây dựng REST API cho máy chủ Node.js."),
    ("Mongoose 8", "ORM kết nối và định nghĩa schema cho MongoDB."),
    ("Socket.IO 4", "Giao tiếp thời gian thực phía server (chat, thông báo)."),
    ("jsonwebtoken", "Sinh và kiểm tra JWT access token/refresh token."),
    ("bcryptjs", "Mã hóa một chiều mật khẩu người dùng."),
    ("express-rate-limit", "Giới hạn tốc độ truy cập API, chống brute-force và quá tải."),
    ("express-validator / joi", "Kiểm tra (validate) dữ liệu đầu vào của API."),
    ("helmet", "Thiết lập các tiêu đề HTTP bảo mật."),
    ("cors", "Quản lý chính sách Cross-Origin Resource Sharing."),
    ("multer", "Xử lý upload file/hình ảnh."),
    ("node-cron", "Lập lịch chạy các tác vụ định kỳ (hủy đơn quá hạn, giải phóng kho...)."),
    ("nodemailer", "Gửi email (xác thực OTP, thông báo)."),
    ("google-auth-library", "Xác thực token đăng nhập Google phía server."),
    ("winston / morgan", "Ghi log hệ thống và log truy cập HTTP."),
], widths=[5, 10.7])

add_table_caption(doc, "5.3")
add_simple_table(doc, ["Thư viện / Công cụ", "Vai trò"], [
    ("FastAPI 0.111", "Framework xây dựng API cho dịch vụ chatbot bằng Python."),
    ("Uvicorn", "ASGI server chạy ứng dụng FastAPI."),
    ("httpx", "Gửi yêu cầu HTTP bất đồng bộ tới API của Server (lấy dữ liệu thật) và tới Ollama."),
    ("python-dotenv", "Quản lý biến môi trường (BOT_MODE, OLLAMA_URL, OLLAMA_MODEL...)."),
    ("Ollama", "Chạy mô hình ngôn ngữ lớn (LLM) đã fine-tune cục bộ, phục vụ sinh câu trả lời."),
], widths=[5, 10.7])

add_heading(doc, "5.2. Tổ chức mã nguồn", level=2)
add_para(doc,
    "Mã nguồn của hệ thống được tổ chức thành ba dự án độc lập (client, server, bot), theo đúng "
    "cách phân chia module đã thiết kế tại Package Diagram (Hình 4.13):", indent_first=1)
add_bullets(doc, [
    "client/src/pages: các trang giao diện (Trang chủ, Sản phẩm, Giỏ hàng, Thanh toán, Hồ sơ, các "
    "trang quản trị...).",
    "client/src/components: các thành phần giao diện dùng lại được (PhoneCard, CartItem, "
    "ChatWidget...).",
    "client/src/hooks, client/src/context: các hook và context quản lý trạng thái dùng chung "
    "(xác thực, giỏ hàng, giao diện, chat realtime...).",
    "client/src/api: các hàm gọi REST API tới Server.",
    "server/src/routes, server/src/controllers: định nghĩa và xử lý các endpoint API theo từng "
    "nhóm nghiệp vụ (auth, product, order, payment, loyalty, wallet, return, chat...).",
    "server/src/models: định nghĩa 23 schema Mongoose tương ứng 23 lớp tại Hình 4.1.",
    "server/src/middlewares: xác thực JWT, kiểm tra phân quyền, giới hạn tốc độ truy cập.",
    "server/src/jobs: các tác vụ định kỳ (cron) tự động hủy đơn quá hạn, giải phóng tồn kho, sinh "
    "mã giảm giá sinh nhật.",
    "bot/main.py, bot/intent.py, bot/actions.py, bot/templates.py, bot/llm.py: lần lượt là điểm "
    "khởi động dịch vụ FastAPI, bộ nhận diện ý định câu hỏi, các hành động truy xuất dữ liệu thật, "
    "các mẫu câu trả lời và lớp giao tiếp với Ollama.",
])

add_heading(doc, "5.3. Một số chức năng nổi bật đã cài đặt", level=2)
add_para(doc, "Cơ chế giữ chỗ và chống bán vượt tồn kho", bold=True, space_after=2)
add_para(doc,
    "Khi áp dụng Flash Sale, số lượng đã bán (sold) được tăng bằng một phép cập nhật atomic có "
    "điều kiện ($expr: sold + quantity <= limit) trực tiếp trên MongoDB, đảm bảo không vượt số "
    "lượng giới hạn của chương trình ngay cả khi nhiều khách hàng đặt hàng cùng lúc, mà không cần "
    "cơ chế khóa (lock) ở mức ứng dụng.", indent_first=1)
add_para(doc, "Tác vụ định kỳ (Cron Jobs)", bold=True, space_after=2)
add_para(doc,
    "Hệ thống chạy các tác vụ định kỳ bằng thư viện node-cron để: tự động hủy các đơn hàng quá hạn "
    "chờ thanh toán, giải phóng tồn kho đã được giữ chỗ trong giỏ hàng nhưng chưa thanh toán, và tự "
    "động sinh mã giảm giá vào ngày sinh nhật của khách hàng.", indent_first=1)
add_para(doc, "Giới hạn tốc độ truy cập (Rate Limiting)", bold=True, space_after=2)
add_para(doc,
    "Ba mức giới hạn được áp dụng bằng thư viện express-rate-limit: tối đa 10 lần đăng nhập trong "
    "15 phút (chống dò mật khẩu), tối đa 5 yêu cầu OTP trong 1 giờ (chống spam email/SMS) và tối đa "
    "100 request/phút cho các API còn lại.", indent_first=1)
add_para(doc, "Chatbot AI tư vấn theo cơ chế RAG", bold=True, space_after=2)
add_para(doc,
    "Dịch vụ chatbot chạy chế độ \"llm\" (cấu hình qua biến môi trường BOT_MODE), sử dụng mô hình "
    "tinh chỉnh riêng \"phonestore-bot\" chạy qua Ollama tại địa chỉ cục bộ. Trước khi trả lời các "
    "câu hỏi liên quan tới giá, tồn kho hoặc thông tin sản phẩm cụ thể, dịch vụ luôn gọi lại API "
    "thật của Server để lấy dữ liệu mới nhất, theo đúng thiết kế RAG đã trình bày tại Hình 4.10.", indent_first=1)

add_heading(doc, "5.4. Giao diện hệ thống", level=2)
add_para(doc,
    "Phần này trình bày các ảnh chụp giao diện thực tế của hệ thống PhoneStore đang chạy, được "
    "chụp trực tiếp trên hệ thống đã cài đặt hoàn chỉnh (không phải bản dựng/mockup), bằng hai tài "
    "khoản thử nghiệm: một tài khoản khách hàng (vai trò user) và một tài khoản quản trị viên (vai "
    "trò admin).", indent_first=1)

add_para(doc, "Trang chủ giới thiệu các sản phẩm nổi bật, sản phẩm bán chạy, sản phẩm mới nhất và "
    "các banner khuyến mãi.", indent_first=1)
add_figure(doc, "5.1")

add_para(doc, "Trang danh sách sản phẩm cho phép khách hàng xem, tìm kiếm và lọc sản phẩm theo "
    "thương hiệu, danh mục.", indent_first=1)
add_figure(doc, "5.2")

add_para(doc, "Trang chi tiết sản phẩm hiển thị đầy đủ thông số kỹ thuật, các biến thể màu sắc/"
    "dung lượng kèm giá và tồn kho tương ứng.", indent_first=1)
add_figure(doc, "5.3")

add_para(doc, "Trang đăng nhập hỗ trợ đăng nhập bằng email/mật khẩu, Google, Facebook, kèm hai "
    "tài khoản thử nghiệm nhanh ở môi trường phát triển.", indent_first=1)
add_figure(doc, "5.4")

add_para(doc, "Trang giỏ hàng cho phép khách hàng xem lại sản phẩm đã chọn, điều chỉnh số lượng và "
    "áp dụng mã giảm giá trước khi thanh toán.", indent_first=1)
add_figure(doc, "5.5")

add_para(doc, "Trang thanh toán cho phép khách hàng nhập địa chỉ giao hàng, chọn phương thức thanh "
    "toán và xác nhận đặt hàng.", indent_first=1)
add_figure(doc, "5.6")

add_para(doc, "Trang hồ sơ cá nhân hiển thị thông tin tài khoản, hạng thành viên hiện tại, điểm "
    "thưởng tích lũy và tiến trình lên hạng tiếp theo.", indent_first=1)
add_figure(doc, "5.7")

add_para(doc, "Trang danh sách đơn hàng cho phép khách hàng theo dõi trạng thái các đơn hàng đã "
    "đặt.", indent_first=1)
add_figure(doc, "5.8")

add_para(doc, "Khung chat hỗ trợ tích hợp chatbot AI: khách hàng đặt câu hỏi và nhận câu trả lời tư "
    "vấn tự động dựa trên dữ liệu sản phẩm thật của hệ thống.", indent_first=1)
add_figure(doc, "5.9")

add_para(doc, "Trang Dashboard quản trị tổng hợp các số liệu kinh doanh: doanh thu, số đơn hàng, "
    "tổng số người dùng, biểu đồ doanh thu theo thời gian và danh sách đơn hàng gần đây.", indent_first=1)
add_figure(doc, "5.10")

add_para(doc, "Trang quản lý sản phẩm cho phép quản trị viên xem, thêm, sửa, xóa sản phẩm và quản "
    "lý tồn kho.", indent_first=1)
add_figure(doc, "5.11")

add_para(doc, "Trang quản lý đơn hàng cho phép quản trị viên xem danh sách và cập nhật trạng thái "
    "xử lý của từng đơn hàng.", indent_first=1)
add_figure(doc, "5.12")

add_heading(doc, "5.5. Kiểm thử", level=2)
add_heading(doc, "5.5.1. Kiểm thử API bằng Postman", level=3)
add_para(doc,
    "Các API của hệ thống được kiểm thử bằng bộ sưu tập Postman (PhoneStore.postman_collection.json) "
    "gồm 31 request, được tổ chức thành 6 nhóm chức năng như trình bày tại Bảng 5.4, bao gồm các "
    "trường hợp xác thực, lấy dữ liệu sản phẩm/danh mục, các API yêu cầu đăng nhập (giỏ hàng, đơn "
    "hàng) và một API kiểm tra tình trạng hoạt động (health check) của máy chủ.", indent_first=1)
add_table_caption(doc, "5.4")
add_simple_table(doc, ["Nhóm API", "Số lượng request", "Ghi chú"], [
    ("Auth", "8", "Đăng ký, đăng nhập, OTP, quên/đặt lại mật khẩu, đăng nhập Google/Facebook."),
    ("Products", "7", "Danh sách, tìm kiếm, chi tiết sản phẩm."),
    ("Brands & Categories", "4", "Danh sách thương hiệu và danh mục."),
    ("Cart (cần đăng nhập)", "7", "Thêm/sửa/xóa sản phẩm trong giỏ, áp mã giảm giá."),
    ("Orders (cần đăng nhập)", "4", "Tạo đơn hàng, xem danh sách/chi tiết đơn, hủy đơn."),
    ("Health Check", "1", "Kiểm tra máy chủ đang hoạt động."),
], widths=[5, 4, 6.7])

add_heading(doc, "5.5.2. Kiểm thử giao diện", level=3)
add_para(doc,
    "Các luồng giao diện chính được kiểm thử thủ công trực tiếp trên hệ thống đang chạy thật, bằng "
    "hai tài khoản thử nghiệm (user@test.com vai trò khách hàng và admin@test.com vai trò quản "
    "trị viên), bao gồm: đăng nhập, xem sản phẩm, thêm vào giỏ hàng, đặt hàng, xem hồ sơ/điểm "
    "thưởng, xem đơn hàng, trò chuyện với chatbot AI, và các chức năng quản trị (dashboard, quản lý "
    "sản phẩm, quản lý đơn hàng). Toàn bộ ảnh chụp ở mục 5.4 được lấy trực tiếp từ các lần kiểm thử "
    "này, không qua dựng giao diện giả lập.", indent_first=1)

add_heading(doc, "5.5.3. Kiểm thử nghiệp vụ đặc biệt", level=3)
add_para(doc,
    "Các nghiệp vụ có khả năng xảy ra tranh chấp dữ liệu khi nhiều người dùng cùng tác động (ví dụ "
    "nhiều khách hàng cùng mua một sản phẩm đang Flash Sale với số lượng giới hạn, hoặc cùng sử "
    "dụng một mã giảm giá có hạn mức) được thiết kế để xử lý bằng các phép cập nhật atomic có điều "
    "kiện trên MongoDB (mục 5.3) ngay từ giai đoạn thiết kế, nhằm tránh tình trạng bán vượt tồn "
    "kho/vượt hạn mức khi có truy cập đồng thời. Trong phạm vi khóa luận, các luồng nghiệp vụ này "
    "đã được kiểm thử thủ công tuần tự để xác nhận đúng logic xử lý (kiểm tra điều kiện, tính toán "
    "số tiền/điểm, cập nhật trạng thái); việc kiểm thử tải với số lượng người dùng đồng thời lớn "
    "trong môi trường sản xuất thực tế chưa được thực hiện và được đề xuất là một hướng phát triển "
    "tiếp theo của đề tài (mục 6.2).", indent_first=1)

# ===================== CHUONG 6 =====================
add_heading(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1, page_break_before=True)

add_heading(doc, "6.1. Kết quả đạt được", level=2)
add_para(doc,
    "Sau quá trình thực hiện, khóa luận đã đạt được các kết quả chính sau:", indent_first=1)
add_bullets(doc, [
    "Áp dụng đầy đủ và có hệ thống phương pháp OOAD/UML theo khung tiến trình lặp và tăng trưởng "
    "(Inception - Elaboration - Construction - Transition) cho toàn bộ quá trình phân tích và "
    "thiết kế hệ thống, thể hiện qua 17 sơ đồ thuộc đủ 8 loại sơ đồ UML được trình bày trong tài "
    "liệu tham khảo [1] (Use Case, Class, Collaboration, Sequence, State, Package, Component, "
    "Deployment) cùng một Activity Diagram bổ sung.",
    "Xây dựng hoàn chỉnh hệ thống thương mại điện tử bán điện thoại PhoneStore với 23 thực thể dữ "
    "liệu, đầy đủ các nghiệp vụ: quản lý sản phẩm theo biến thể, giỏ hàng, đặt hàng, thanh toán "
    "(COD/ví điện tử/VNPay), Flash Sale, mã giảm giá, điểm thưởng theo hạng thành viên, đánh giá "
    "sản phẩm, wishlist, đổi trả hàng và một trang quản trị đầy đủ chức năng.",
    "Tích hợp thành công một chatbot tư vấn bằng AI sử dụng mô hình ngôn ngữ lớn đã fine-tune riêng "
    "cho lĩnh vực bán điện thoại, hoạt động theo cơ chế RAG để đảm bảo câu trả lời về giá/tồn kho "
    "luôn dựa trên dữ liệu thật của hệ thống tại thời điểm hỏi.",
    "Cài đặt và kiểm thử thành công hệ thống trên môi trường thực tế, với bộ kiểm thử API gồm 31 "
    "request bằng Postman và kiểm thử giao diện trực tiếp trên toàn bộ các luồng nghiệp vụ chính.",
])

add_heading(doc, "6.2. Hạn chế", level=2)
add_bullets(doc, [
    "Hệ thống chưa được kiểm thử tải (load test) với số lượng người dùng truy cập đồng thời lớn "
    "trong môi trường sản xuất thực tế.",
    "Cổng thanh toán VNPay mới được tích hợp ở môi trường thử nghiệm (sandbox), chưa xử lý thanh "
    "toán thật.",
    "Chatbot AI chạy mô hình ngôn ngữ lớn cục bộ qua Ollama nên tốc độ phản hồi và khả năng mở "
    "rộng phụ thuộc trực tiếp vào tài nguyên phần cứng của máy chủ AI, chưa được kiểm thử trên hạ "
    "tầng GPU/cloud chuyên dụng.",
    "Đề tài chưa xây dựng ứng dụng di động (native mobile app) riêng cho hệ điều hành Android/iOS.",
])

add_heading(doc, "6.3. Hướng phát triển", level=2)
add_bullets(doc, [
    "Thực hiện kiểm thử tải (ví dụ bằng k6 hoặc JMeter) để đánh giá khả năng chịu tải thực tế của "
    "hệ thống, đặc biệt với các nghiệp vụ có tranh chấp dữ liệu cao như Flash Sale.",
    "Container hóa hệ thống bằng Docker và xây dựng pipeline CI/CD để triển khai tự động.",
    "Tích hợp thanh toán VNPay ở môi trường thật và bổ sung thêm các cổng thanh toán khác (Momo, "
    "ZaloPay...).",
    "Xây dựng ứng dụng di động đa nền tảng (React Native/Flutter) để mở rộng kênh tiếp cận khách "
    "hàng.",
    "Mở rộng chatbot AI hỗ trợ đa ngôn ngữ và bổ sung mô hình gợi ý sản phẩm (recommendation "
    "system) dựa trên lịch sử mua hàng và hành vi của khách hàng.",
])

# ===================== TAI LIEU THAM KHAO =====================
add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1, page_break_before=True)
refs = [
    "[1] UML Applied - Object Oriented Analysis and Design using the UML (giáo trình/tài liệu "
    "tham khảo, 123 trang) - được sử dụng làm nền tảng phương pháp luận OOAD/UML và quy trình phát "
    "triển RUP áp dụng xuyên suốt khóa luận (Chương 2, 3, 4).",
    "[2] Express.js Official Documentation - https://expressjs.com - tài liệu chính thức về "
    "framework Express.js dùng để xây dựng REST API phía server.",
    "[3] React Official Documentation - https://react.dev - tài liệu chính thức về thư viện React "
    "dùng để xây dựng giao diện người dùng phía client.",
    "[4] MongoDB Documentation - https://www.mongodb.com/docs - tài liệu chính thức về cơ sở dữ "
    "liệu NoSQL MongoDB và thư viện Mongoose.",
    "[5] FastAPI Documentation - https://fastapi.tiangolo.com - tài liệu chính thức về framework "
    "FastAPI dùng để xây dựng dịch vụ chatbot AI.",
    "[6] Socket.IO Documentation - https://socket.io/docs - tài liệu chính thức về thư viện "
    "Socket.IO dùng cho giao tiếp thời gian thực (chat, thông báo).",
]
for r in refs:
    add_para(doc, r, align='justify', space_after=8)

doc.save(OUT)
print("HOAN TAT - Da luu file day du tai:", OUT)
